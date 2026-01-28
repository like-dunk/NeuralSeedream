"""
图文裂变服务模块
基于 OpenRouter/OpenAI 模型生成多个语义一致但表达不同的文案变体
支持批量生成和并发控制
"""
import json
import logging
import math
import re
import uuid
import asyncio
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Optional, Any

from openai import AsyncOpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement

from reflo_dashboard.core.config import settings

logger = logging.getLogger(__name__)


# ==================== Few-shot 样本库 ====================
FEW_SHOT_EXAMPLES = [
    {
        "original": {
            "title": "谁懂啊！上班摸鱼满脑子都是武网的挥拍声😭",
            "content": "谁能理解这种痛苦啊！武网的比赛场面还在我脑海里挥之不去，人已经坐在办公室对着电脑敲敲打打，这落差感真的要把我击垮了！现在只要一闭眼，满脑子都是选手们在赛场上潇洒挥拍的画面，而且我想说作为赞助商的岚图汽车也太会了，不仅有超酷的展车摆在赛场边，看比赛的时候我都忍不住频频瞟过去，离场还拿到了周边，诚意直接拉满！真的好想穿越回假期啊！能让我再回到武网现场吗 555～打工人真的需要再来一场比赛续命！"
        },
        "variation": {
            "title": "破防了！一上班就想起武网现场，根本静不下心😭",
            "content": "真的要疯了！假期在武网看比赛的记忆还历历在目，现在却要面对电脑码字，这种心理落差谁能懂啊！每次闭眼脑海里就浮现赛场上的精彩瞬间，还有岚图汽车在现场的展车，设计感绝了，当时看比赛的时候眼神都不自觉往那边飘。离场还送了周边，真的太贴心了！现在好想时光倒流回到武网啊 QAQ 打工人急需下一场比赛来回血！"
        }
    },
    {
        "original": {
            "title": "桂林山水双车同框！这颜值也太绝了",
            "content": "桂林山水竟是写实画！看看车和水墨山水到底有多适配～追光L的车漆在阳光下就像清晨漓江水的微光，车身线条流畅大气，前脸的镀铬饰条稳重精致；享界S9线条更圆润，流线型车身更偏年轻运动一些。同样的背景下两款车哪一款更出色呢？享界S9售价30万左右，追光L还没公布价格，什么时候上市大家有消息不？但听说最近追光L展车到店了，有兴趣的宝子可以去看看～"
        },
        "variation": {
            "title": "桂林实拍！两款豪车与山水同框美翻了",
            "content": "没想到桂林山水真的像画一样！这次拍到了车与山水的完美融合～追光L在阳光映照下，车漆泛着漓江晨光般的光泽，整体线条大气流畅，镀铬前脸显得格外精致；享界S9则更圆润，流线车身充满年轻动感。在这样的背景下，你觉得哪台车更惊艳？享界S9大概30万价位，追光L的定价还没出，上市时间有人知道吗？不过最近追光L展车已经到店了，感兴趣的朋友可以去实地看看～"
        }
    }
]


class TextFissionService:
    """图文裂变服务类 - 基于 OpenRouter/OpenAI 模型，支持批量生成和并发控制"""
    
    # 并发控制常量
    BATCH_SIZE = 5       # 每次 API 请求生成的变体数量
    MAX_CONCURRENT = 10  # 最大并发请求数
    MAX_VARIATIONS = 200 # 单次裂变最大变体数量限制
    
    def __init__(self):
        """初始化服务"""
        self.api_key = settings.OPENROUTER_API_KEY
        self.base_url = settings.OPENROUTER_BASE_URL
        self.site_url = settings.OPENROUTER_SITE_URL
        self.site_name = settings.OPENROUTER_SITE_NAME
        self.model = settings.OPENAI_MODEL
        self.temperature = settings.OPENAI_TEMPERATURE
        self.max_retries = settings.OPENAI_MAX_RETRIES
        self.default_count = settings.TEXT_FISSION_DEFAULT_COUNT
        
        # 初始化 OpenAI 异步客户端
        self.openai_client: Optional[AsyncOpenAI] = None
        if self.api_key:
            self.openai_client = AsyncOpenAI(
                api_key=self.api_key,
                base_url=self.base_url
            )
    
    def is_enabled(self) -> bool:
        """检查服务是否启用（API Key 已配置）"""
        return bool(self.api_key) and self.openai_client is not None
    
    def _build_system_prompt(self) -> str:
        """构建系统提示词 - 批量生成5个变体"""
        return f"""你是一个专业的营销文案改写专家。你的任务是基于原始标题和文案，一次性生成 {self.BATCH_SIZE} 个语义一致但表达各不相同的新版本。

【⚠️⚠️⚠️ 最重要的格式要求 - 必须严格遵守 ⚠️⚠️⚠️】
你**必须且只能**返回一个有效的JSON对象，包含 {self.BATCH_SIZE} 个变体的数组，格式如下：
{{
  "variations": [
    {{"title": "改写后的标题1", "content": "改写后的文案1"}},
    {{"title": "改写后的标题2", "content": "改写后的文案2"}},
    {{"title": "改写后的标题3", "content": "改写后的文案3"}},
    {{"title": "改写后的标题4", "content": "改写后的文案4"}},
    {{"title": "改写后的标题5", "content": "改写后的文案5"}}
  ]
}}

**严格禁止以下行为：**
❌ 在JSON前后添加任何文字说明、解释或注释
❌ 使用代码块标记（如 ```json 或 ```）
❌ 返回不完整的JSON（被截断或缺少字段）
❌ 在JSON中添加注释或说明文字
❌ 返回少于 {self.BATCH_SIZE} 个变体

**必须遵守：**
✅ 只返回纯JSON，可以直接被 json.loads() 解析
✅ JSON必须包含 variations 数组，数组中必须有 {self.BATCH_SIZE} 个对象
✅ 每个对象必须包含 title 和 content 两个字段，且都是字符串类型
✅ 确保所有双引号正确闭合
✅ 如果内容较长，必须完整输出，不能省略或截断
✅ {self.BATCH_SIZE} 个变体之间要有足够的差异化，避免雷同

【核心要求】
1. 保持语义一致：不改变营销重点、产品特性、情感基调
2. 禁止虚构信息：不添加原文未提及的任何信息（产品、品牌、数据等）
3. 表达多样化：使用不同的句式、词汇、语气，但保持自然流畅
4. 保留关键信息：品牌名、产品名、价格、时间等必须保持一致
5. 情感连贯：保持原文的情感强度和表达风格（如：激动、惊喜、疑问等）
6. 变体差异化：{self.BATCH_SIZE} 个变体之间要有明显区别，不能只是简单换词

【标题改写规则 - 极其重要！】
⚠️ 标题改写必须严格遵循以下原则：
1. **标题只能基于原始标题改写**，绝对不能根据改写后的内容重新提炼标题
2. **禁止总结式标题**：不要把改写后的内容压缩成标题
3. **保持标题结构相似**：如果原标题是对比式，改写标题也应该是对比式
4. **改写方式**：
   - 调整词汇：如"设计对决"→"设计PK"→"设计较量"→"设计之争"
   - 调整语气：如"VS"→"对比"→"对决"→"较量"
   - 调整顺序：如"A VS B：C"→"C：A与B的较量"
   - 调整修饰：如"东方美学"→"中式美学"→"东方韵味"→"中式格调"
5. **核心主题不变**：原标题讲什么，改写标题必须讲什么

【内容改写规则】
1. **内容只能基于原始内容改写**，保持原文的信息点和逻辑结构
2. 可以调整句式、词汇、表达顺序，但不能删减关键信息
3. 不能根据标题去重新组织内容

【禁止行为】
- ❌ 不要根据改写后的内容重新提炼标题
- ❌ 不要编造原文没有的品牌、产品、功能、价格
- ❌ 不要改变原文的核心观点和营销目的
- ❌ 不要过度夸张或削弱情感表达
- ❌ 不要改变数字、日期、专有名词
- ❌ 不要让 {self.BATCH_SIZE} 个变体太过相似"""

    def _build_few_shot_prompt(self) -> str:
        """构建 Few-shot 示例提示词 - 展示批量输出格式"""
        # 使用第一个示例展示批量输出格式
        ex = FEW_SHOT_EXAMPLES[0]
        example_output = {
            "variations": [
                ex['variation'],
                {
                    "title": "上班第一天就开始怀念武网了，这种感觉谁懂😭",
                    "content": "假期余韵还没散去呢！武网看比赛的画面一直在脑海里循环播放，现在对着电脑屏幕完全提不起劲。闭上眼就是赛场上选手们精彩的挥拍瞬间，对了还有岚图汽车的展车，当时在现场看得移不开眼，设计真的太酷了，走的时候还领了周边回来。好想立刻回到武网现场啊！打工人太需要比赛来续命了呜呜呜！"
                },
                {
                    "title": "才上班就满脑子武网画面，心思完全不在工作上😭",
                    "content": "这心理落差真的绷不住了！武网的精彩比赛还历历在目，人却已经坐在工位上开始码字。每次放空脑海里就是选手挥拍的飒爽身影，还有岚图汽车在现场摆的展车，颜值设计感都在线，当时看比赛眼神总是飘过去。临走还拿了周边，太有诚意了！现在只想穿越回武网现场 QAQ 打工人真的需要靠比赛续命啊！"
                },
                {
                    "title": "开工第一天脑子里全是武网，完全静不下来😭",
                    "content": "谁懂这种痛苦！武网看比赛的快乐还在回味中，结果人已经回到办公室敲键盘了。只要一发呆脑海里就是赛场上的精彩画面，说起来岚图汽车在现场的展示太吸睛了，展车设计超有范儿，当时忍不住多看了好几眼。而且离场还发了周边，真的很贴心！好想再去一次武网啊555 打工人急需再来一场比赛充充电！"
                },
                {
                    "title": "工作日开始了但心还留在武网，谁来救救我😭",
                    "content": "太难熬了！假期在武网的美好记忆挥之不去，现在坐在电脑前根本集中不了注意力。一闭眼全是比赛时选手潇洒挥拍的场景，还有岚图汽车摆在现场的展车，设计太赞了，看比赛时目光总被吸引过去。走的时候还收到了周边礼品，满满的诚意！现在好想穿越回去啊！打工人真的太需要下一场比赛来回血了！"
                }
            ]
        }
        
        return f"""以下是一个改写示例，展示如何一次生成 {self.BATCH_SIZE} 个不同的变体。

⚠️ 注意：输出格式必须是包含 variations 数组的 JSON 对象，数组中有 {self.BATCH_SIZE} 个变体。

【示例】
输入（原标题和原文案）：
标题：{ex['original']['title']}
文案：{ex['original']['content']}

输出（必须严格按照以下JSON格式，包含 {self.BATCH_SIZE} 个变体）：
{json.dumps(example_output, ensure_ascii=False, indent=2)}"""
    
    def _build_user_prompt(self, title: str, content: str) -> str:
        """构建用户提示词 - 要求一次生成5个变体"""
        return f"""请基于以下原始内容，一次性生成 {self.BATCH_SIZE} 个不同的改写版本。

⚠️ 重要提醒：
- 标题改写：只能基于【原标题】进行改写，不能根据改写后的内容重新提炼
- 内容改写：只能基于【原文案】进行改写
- {self.BATCH_SIZE} 个变体之间要有明显差异，不能只是简单换几个词

原标题（改写标题时只能基于这个）：
{title}

原文案（改写文案时只能基于这个）：
{content}

【⚠️⚠️⚠️ 输出格式要求 - 必须严格遵守 ⚠️⚠️⚠️】
你必须**只返回一个JSON对象**，包含 {self.BATCH_SIZE} 个变体的数组：

{{
  "variations": [
    {{"title": "改写版本1的标题", "content": "改写版本1的文案"}},
    {{"title": "改写版本2的标题", "content": "改写版本2的文案"}},
    {{"title": "改写版本3的标题", "content": "改写版本3的文案"}},
    {{"title": "改写版本4的标题", "content": "改写版本4的文案"}},
    {{"title": "改写版本5的标题", "content": "改写版本5的文案"}}
  ]
}}

**严格禁止：**
❌ 不要使用 ```json 或 ``` 代码块标记
❌ 不要在JSON前后添加说明文字
❌ 不要返回少于 {self.BATCH_SIZE} 个变体
❌ 不要让变体之间太过相似

**必须遵守：**
✅ 只返回纯JSON，格式必须可以直接被 json.loads() 解析
✅ variations 数组必须包含 {self.BATCH_SIZE} 个对象
✅ 每个对象必须包含 title 和 content 字段
✅ 确保JSON格式完整，不要被截断
✅ {self.BATCH_SIZE} 个变体要有足够的差异化"""
    
    def _extract_json(self, text: str) -> Optional[Dict]:
        """
        从文本中提取JSON内容
        
        Args:
            text: 可能包含JSON的文本
            
        Returns:
            解析后的字典，失败返回None
        """
        def parse_candidate(json_text: str, method_label: str, *, show_snippet: bool = False) -> Optional[Dict]:
            """尝试解析JSON"""
            try:
                result = json.loads(json_text)
                if isinstance(result, dict):
                    logger.debug(f"[JSON解析] {method_label}成功: 直接解析")
                    return result
            except json.JSONDecodeError as e:
                logger.debug(f"[JSON解析] {method_label}失败: {str(e)}")
                if "control character" in str(e):
                    try:
                        result = json.loads(json_text, strict=False)
                        if isinstance(result, dict):
                            logger.debug(f"[JSON解析] {method_label}成功: 使用 strict=False")
                            return result
                    except:
                        pass
            except Exception as e:
                logger.debug(f"[JSON解析] {method_label}异常: {type(e).__name__}: {str(e)}")
            return None
        
        # 尝试1: 直接解析
        result = parse_candidate(text, "方法1")
        if result:
            return result
        
        # 尝试2: 提取代码块中的JSON
        json_pattern = r'```(?:json)?\s*(\{.*?\})\s*```'
        match = re.search(json_pattern, text, re.DOTALL)
        if match:
            result = parse_candidate(match.group(1), "方法2")
            if result:
                return result
        
        # 尝试3: 提取完整的大括号内容
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1 and end > start:
            json_str = text[start:end+1]
            result = parse_candidate(json_str, "方法3")
            if result:
                return result
        
        # 尝试4: 处理被截断的JSON
        if start != -1:
            json_str = text[start:]
            if '"content"' in json_str and not json_str.rstrip().endswith('}'):
                json_str = json_str.rstrip()
                if not json_str.endswith('"'):
                    json_str += '"'
                if not json_str.endswith('}'):
                    json_str += '\n}'
                result = parse_candidate(json_str, "方法4")
                if result and 'title' in result and 'content' in result:
                    return result
        
        logger.warning(f"[JSON解析] 所有方法均失败，原始文本前200字符: {text[:200]}")
        return None
    
    async def _call_openai_api(self, title: str, content: str) -> List[Dict]:
        """
        调用 OpenRouter/OpenAI API 批量生成裂变版本
        
        Args:
            title: 原始标题
            content: 原始文案
            
        Returns:
            包含多个变体的列表，每个变体包含 title 和 content，失败返回空列表
        """
        if not self.openai_client:
            logger.error("[OpenAI API] 客户端未初始化")
            return []
        
        # 构建提示词
        system_prompt = self._build_system_prompt()
        few_shot_prompt = self._build_few_shot_prompt()
        user_prompt = self._build_user_prompt(title, content)
        
        # 组合系统提示词
        full_system_prompt = f"{system_prompt}\n\n{few_shot_prompt}"
        
        # 重试机制
        for attempt in range(self.max_retries):
            try:
                logger.info(f"[OpenAI API] 尝试 {attempt + 1}/{self.max_retries}，请求生成 {self.BATCH_SIZE} 个变体")
                
                # 构建 extra_headers
                extra_headers = {}
                if self.site_url:
                    extra_headers["HTTP-Referer"] = self.site_url
                if self.site_name:
                    extra_headers["X-Title"] = self.site_name
                
                completion = await self.openai_client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": full_system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    extra_headers=extra_headers if extra_headers else None,
                    response_format={"type": "json_object"},
                    temperature=self.temperature
                )
                
                if completion.choices and completion.choices[0].message.content:
                    message_content = completion.choices[0].message.content
                    
                    result = self._extract_json(message_content)
                    if result and "variations" in result and isinstance(result["variations"], list):
                        variations = result["variations"]
                        # 验证每个变体都有 title 和 content
                        valid_variations = [
                            v for v in variations 
                            if isinstance(v, dict) and "title" in v and "content" in v
                        ]
                        if valid_variations:
                            logger.info(f"[OpenAI API] 调用成功，获取到 {len(valid_variations)} 个有效变体")
                            return valid_variations
                        else:
                            logger.warning("[OpenAI API] variations 数组中没有有效的变体")
                    else:
                        logger.warning("[OpenAI API] JSON解析失败或缺少 variations 字段")
                else:
                    logger.warning("[OpenAI API] 返回内容为空")
                
            except asyncio.TimeoutError:
                logger.error("[OpenAI API] 请求超时")
            except Exception as e:
                logger.error(f"[OpenAI API] 错误: {type(e).__name__}: {str(e)}")
            
            # 等待后重试
            if attempt < self.max_retries - 1:
                wait_time = (attempt + 1) * 2
                logger.info(f"[OpenAI API] {wait_time}秒后重试...")
                await asyncio.sleep(wait_time)
        
        logger.error(f"[OpenAI API] 已达到最大重试次数 ({self.max_retries})，放弃本次调用")
        return []
    
    async def generate_variations(
        self, 
        title: str, 
        content: str, 
        count: Optional[int] = None
    ) -> List[Dict]:
        """
        生成多个裂变版本（支持并发批量请求）
        
        每次 API 请求生成 BATCH_SIZE 个变体，最多 MAX_CONCURRENT 个并发请求。
        使用信号量实现排队机制，确保不超过最大并发数。
        
        Args:
            title: 原始标题
            content: 原始文案
            count: 裂变数量，默认使用配置中的值
            
        Returns:
            裂变版本列表，每个元素包含 title 和 content
        """
        count = count or self.default_count
        
        # 限制最大变体数量
        if count > self.MAX_VARIATIONS:
            logger.warning(f"请求的变体数量 {count} 超过最大限制 {self.MAX_VARIATIONS}，已自动调整")
            count = self.MAX_VARIATIONS
        
        # 计算需要多少次 API 调用
        batch_count = math.ceil(count / self.BATCH_SIZE)
        
        logger.info(f"开始生成 {count} 个图文裂变版本，需要 {batch_count} 次 API 调用，"
                   f"每次生成 {self.BATCH_SIZE} 个，最大并发 {self.MAX_CONCURRENT}")
        
        # 使用信号量限制并发数
        semaphore = asyncio.Semaphore(self.MAX_CONCURRENT)
        
        async def fetch_batch(batch_num: int) -> List[Dict]:
            """获取单个批次的变体，受信号量控制"""
            async with semaphore:
                logger.info(f"[批次 {batch_num + 1}/{batch_count}] 开始请求...")
                result = await self._call_openai_api(title, content)
                logger.info(f"[批次 {batch_num + 1}/{batch_count}] 完成，获取 {len(result)} 个变体")
                return result
        
        # 创建所有批次的任务
        tasks = [fetch_batch(i) for i in range(batch_count)]
        
        # 并发执行所有任务（信号量会自动控制并发数）
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 合并所有结果
        variations = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"[批次 {i + 1}] 执行失败: {type(result).__name__}: {str(result)}")
            elif isinstance(result, list):
                variations.extend(result)
            else:
                logger.warning(f"[批次 {i + 1}] 返回了意外的类型: {type(result)}")
        
        # 截取到所需数量
        if len(variations) > count:
            variations = variations[:count]
        
        if len(variations) < count:
            logger.warning(f"仅成功生成 {len(variations)}/{count} 个版本")
        else:
            logger.info(f"成功生成 {len(variations)}/{count} 个版本")
        
        return variations
    
    def _set_run_font(self, run):
        """为 run 设置字体（包括中文字体支持）"""
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Arial')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Arial')
    
    def create_word_document(
        self,
        theme: str,
        original_title: str,
        original_content: str,
        tags: str,
        variations: List[Dict]
    ) -> BytesIO:
        """
        创建 Word 文档
        
        Args:
            theme: 主题名称
            original_title: 原始标题
            original_content: 原始文案
            tags: 标签
            variations: 裂变版本列表
        
        Returns:
            BytesIO 对象包含 Word 文档内容
        """
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 
            '微软雅黑'
        )
        
        # 添加标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"主题：{theme}")
        title_run.bold = True
        title_run.font.size = Pt(16)
        self._set_run_font(title_run)
        
        doc.add_paragraph()
        
        # 添加原始版本
        original_title_para = doc.add_paragraph()
        original_title_run = original_title_para.add_run("原始版本：")
        original_title_run.bold = True
        original_title_run.font.size = Pt(14)
        self._set_run_font(original_title_run)
        
        # 原始标题
        p = doc.add_paragraph()
        run = p.add_run("标题：")
        run.bold = True
        self._set_run_font(run)
        run = p.add_run(original_title)
        self._set_run_font(run)
        
        # 原始文案
        p = doc.add_paragraph()
        run = p.add_run("文案：")
        run.bold = True
        self._set_run_font(run)
        run = p.add_run(original_content)
        self._set_run_font(run)
        
        # 标签
        if tags:
            p = doc.add_paragraph()
            run = p.add_run(tags)
            self._set_run_font(run)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # 添加各个裂变版本
        for i, variation in enumerate(variations, 1):
            version_title_para = doc.add_paragraph()
            version_run = version_title_para.add_run(f"版本{i}：")
            version_run.bold = True
            version_run.font.size = Pt(14)
            self._set_run_font(version_run)
            
            # 版本标题
            p = doc.add_paragraph()
            run = p.add_run("标题：")
            run.bold = True
            self._set_run_font(run)
            run = p.add_run(variation.get("title", ""))
            self._set_run_font(run)
            
            # 版本文案
            p = doc.add_paragraph()
            run = p.add_run("文案：")
            run.bold = True
            self._set_run_font(run)
            run = p.add_run(variation.get("content", ""))
            self._set_run_font(run)
            
            # 标签
            if tags:
                p = doc.add_paragraph()
                run = p.add_run(tags)
                self._set_run_font(run)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
        
        # 保存到 BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    async def upload_to_oss(self, file_buffer: BytesIO, filename: str) -> Optional[str]:
        """
        上传 Word 文档到 OSS
        
        Args:
            file_buffer: Word 文档的 BytesIO 对象
            filename: 文件名
            
        Returns:
            OSS 对象键名，失败返回 None
        """
        try:
            from .oss_service import get_oss_service
            oss_service = get_oss_service()
            
            # 构建 OSS 路径
            date_str = datetime.now().strftime("%Y%m%d")
            object_key = f"Reflo-Dashboard/text_fission/{date_str}/{filename}"
            
            # 上传文件
            file_buffer.seek(0)
            result = oss_service.bucket.put_object(object_key, file_buffer.read())
            
            if result.status == 200:
                logger.info(f"Word 文档上传成功: {object_key}")
                return object_key
            else:
                logger.error(f"Word 文档上传失败: status={result.status}")
                return None
                
        except Exception as e:
            logger.error(f"上传 Word 文档到 OSS 失败: {str(e)}", exc_info=True)
            return None
    
    def calculate_fission_count(self, frequency_text: str) -> int:
        """
        根据频次文本计算图文裂变数量
        
        规则：取最大频次值，除以 3，再减 1（图文用除3，视频用除4）
        公式：max_frequency / 3 - 1
        
        Args:
            frequency_text: 频次文本（可能包含多行）
            
        Returns:
            裂变数量（最小为 1）
        """
        if not frequency_text:
            logger.warning("频次文本为空，使用默认裂变数量")
            return self.default_count
        
        try:
            numbers = re.findall(r'\d+', frequency_text)
            
            if not numbers:
                logger.warning(f"频次文本中未找到数字: {frequency_text}")
                return self.default_count
            
            max_frequency = max(int(n) for n in numbers)
            # 图文裂变使用除3的公式
            fission_count = max_frequency // 3 - 1
            fission_count = max(1, fission_count)
            
            logger.info(f"根据频次计算图文裂变数量: max_frequency={max_frequency}, fission_count={fission_count}")
            return fission_count
            
        except Exception as e:
            logger.error(f"计算裂变数量失败: {e}")
            return self.default_count
    
    async def process_text_fission(
        self,
        record_id: int,
        theme: str,
        title: str,
        content: str,
        frequency: Optional[str] = None,
        tags: Optional[str] = None,
        num_fission: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        处理图文裂变的完整流程
        
        Args:
            record_id: 记录ID
            theme: 主题名称
            title: 原始标题（从 content 中提取或使用 theme）
            content: 原始文案
            frequency: 频次文本
            tags: 标签
            num_fission: 前端传入的裂变数量
            
        Returns:
            包含 task_id, variations, oss_key 的结果字典
        """
        task_id = str(uuid.uuid4())
        
        logger.info(f"开始处理图文裂变: record_id={record_id}, task_id={task_id}")
        
        # 使用前端传入的裂变数量，仅在未传时才兜底计算
        fission_count = num_fission if num_fission and num_fission >= 1 else (self.calculate_fission_count(frequency) if frequency else self.default_count)
        
        # 生成裂变版本
        variations = await self.generate_variations(title, content, fission_count)
        
        if not variations:
            logger.error(f"图文裂变失败: 未生成任何变体")
            return {
                "success": False,
                "task_id": task_id,
                "message": "未生成任何文案变体"
            }
        
        # 生成 Word 文档
        doc_buffer = self.create_word_document(
            theme=theme,
            original_title=title,
            original_content=content,
            tags=tags or "",
            variations=variations
        )
        
        # 上传到 OSS
        filename = f"{theme}_{record_id}_{task_id[:8]}.docx"
        oss_key = await self.upload_to_oss(doc_buffer, filename)
        
        if not oss_key:
            logger.error(f"图文裂变失败: Word 文档上传失败")
            return {
                "success": False,
                "task_id": task_id,
                "variations": variations,
                "message": "Word 文档上传失败"
            }
        
        logger.info(f"图文裂变成功: record_id={record_id}, variations={len(variations)}, oss_key={oss_key}")
        
        return {
            "success": True,
            "task_id": task_id,
            "variations": variations,
            "oss_key": oss_key,
            "variation_count": len(variations)
        }


# 全局服务实例（单例模式）
_text_fission_service: Optional[TextFissionService] = None


def get_text_fission_service() -> TextFissionService:
    """获取图文裂变服务实例"""
    global _text_fission_service
    if _text_fission_service is None:
        _text_fission_service = TextFissionService()
    return _text_fission_service

